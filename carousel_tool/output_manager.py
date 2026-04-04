"""
Saves carousel output to local folder structure:
  output/<YYYY-MM-DD>/<article-title>/
    ├── slide_1.png … slide_N.png (generated images)
    ├── slide_1.json … slide_N.json (image prompts)
    ├── social_copy_brief.docx
    ├── carousel_brief.txt (for Instagram caption)
    └── raw_response.txt
"""

import json
import logging
import re
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

from config import Config
from gemini_image_gen import generate_carousel_images

logger = logging.getLogger(__name__)


def sanitize_filename(name: str) -> str:
    """Remove characters that are invalid in filenames."""
    name = re.sub(r'[<>:"/\\|?*\x00-\x1f]', "", name)
    name = name.strip(". ")
    return name[:80] or "untitled"


def save_output(carousel_data: dict) -> Path | None:
    """
    Save all carousel output files to a local folder.
    Returns the folder Path, or None on failure.
    """
    try:
        article = carousel_data["article"]
        title = sanitize_filename(article.get("title", "untitled"))
        date_str = datetime.today().strftime("%Y-%m-%d")

        folder = Config.output_dir() / date_str / title
        folder.mkdir(parents=True, exist_ok=True)

        # ── JSON files ────────────────────────────────────────────────────
        jsons = carousel_data.get("jsons", [])
        for i, json_data in enumerate(jsons, 1):
            dest = folder / f"slide_{i}.json"
            dest.write_text(
                json.dumps(json_data, ensure_ascii=False, indent=2),
                encoding="utf-8",
            )

        # For any missing JSON slots, save slide text as minimal JSON
        slides = carousel_data.get("slides", [])
        for slide in slides:
            dest = folder / f"slide_{slide['number']}.json"
            if not dest.exists():
                dest.write_text(
                    json.dumps(slide, ensure_ascii=False, indent=2),
                    encoding="utf-8",
                )

        # ── Generate carousel images with Gemini ──────────────────────────
        logger.info("Generating carousel images with Gemini...")
        image_result = generate_carousel_images(jsons, folder)
        if image_result["success"]:
            logger.info(
                "✓ Generated %d/%d slide images",
                image_result["generated"],
                image_result["total"],
            )
        else:
            logger.warning(
                "⚠ Failed to generate %d slides: %s",
                len(image_result["failed_slides"]),
                image_result["failed_slides"],
            )

        # ── DOCX social brief ─────────────────────────────────────────────
        docx_path = folder / "social_copy_brief.docx"
        _create_docx(carousel_data, docx_path)

        # ── Instagram caption brief ───────────────────────────────────────
        brief_path = folder / "carousel_brief.txt"
        _create_carousel_brief(carousel_data, brief_path)

        # ── Raw response (debug) ──────────────────────────────────────────
        raw_path = folder / "raw_response.txt"
        raw_path.write_text(
            carousel_data.get("raw_content", ""), encoding="utf-8"
        )

        logger.info("Output saved to: %s", folder)
        return folder

    except Exception as e:
        logger.exception("Failed to save output: %s", e)
        return None


# ── DOCX builder ──────────────────────────────────────────────────────────────

def _set_rtl(paragraph):
    """Apply RTL bidi mark to a paragraph."""
    pPr = paragraph._p.get_or_add_pPr()
    bidi = OxmlElement("w:bidi")
    bidi.set(qn("w:val"), "1")
    pPr.append(bidi)


def _heading(doc: Document, text: str, level: int = 1):
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    _set_rtl(p)
    return p


def _para(doc: Document, text: str, bold_prefix: str = ""):
    p = doc.add_paragraph()
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
    p.add_run(text)
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    _set_rtl(p)
    return p


def _create_docx(carousel_data: dict, path: Path):
    doc = Document()

    # Global RTL for all sections
    for section in doc.sections:
        sectPr = section._sectPr
        bidi = OxmlElement("w:bidi")
        bidi.set(qn("w:val"), "1")
        sectPr.append(bidi)

    article = carousel_data["article"]
    slides = carousel_data.get("slides", [])
    social_brief = carousel_data.get("social_brief", "")

    # ── Header ────────────────────────────────────────────────────────────
    _heading(doc, f"תקציר סושיאל — {article.get('title', '')}", level=1)
    doc.add_paragraph()

    _para(doc, article.get("url", ""), bold_prefix="כתובת כתבה: ")
    _para(doc, article.get("date", ""), bold_prefix="תאריך פרסום: ")
    doc.add_paragraph()

    # ── Slides summary ────────────────────────────────────────────────────
    _heading(doc, "תקציר שקפים", level=2)
    for slide in slides:
        p = doc.add_paragraph(style="List Bullet")
        run_num = p.add_run(f"שקף {slide['number']}: ")
        run_num.bold = True
        p.add_run(slide["title"])
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        _set_rtl(p)

    doc.add_paragraph()

    # ── Full slide texts ──────────────────────────────────────────────────
    _heading(doc, "טקסט מלא לכל שקף", level=2)
    for slide in slides:
        _heading(doc, f"שקף {slide['number']}: {slide['title']}", level=3)
        if slide.get("body"):
            for line in slide["body"].split("\n"):
                if line.strip():
                    _para(doc, line.strip())
        doc.add_paragraph()

    # ── Social copy brief ─────────────────────────────────────────────────
    if social_brief:
        _heading(doc, "הנחיות פרסום סושיאל", level=2)
        for line in social_brief.split("\n"):
            if line.strip():
                _para(doc, line.strip())

    doc.save(str(path))
    logger.info("DOCX brief saved: %s", path)


# ── Instagram caption brief ───────────────────────────────────────────────────

def _create_carousel_brief(carousel_data: dict, path: Path):
    """
    Create a simple text file for Instagram caption.
    User can copy-paste this as the post description.
    """
    try:
        slides = carousel_data.get("slides", [])
        social_brief = carousel_data.get("social_brief", "")
        article = carousel_data["article"]

        lines = []
        lines.append("=== קרוסלה סושיאל ===\n")
        lines.append(f"כתבה: {article.get('title', '')}")
        lines.append(f"תאריך: {article.get('date', '')}")
        lines.append(f"קישור: {article.get('url', '')}\n")

        lines.append("--- טקסט כל שקף ---\n")
        for slide in slides:
            lines.append(f"שקף {slide['number']}:")
            lines.append(f"  כותרת: {slide.get('title', '')}")
            lines.append(f"  טקסט: {slide.get('body', '')}\n")

        lines.append("--- הנחיות פרסום ---\n")
        if social_brief:
            lines.append(social_brief)

        content = "\n".join(lines)
        path.write_text(content, encoding="utf-8")
        logger.info("Carousel brief saved: %s", path)

    except Exception as e:
        logger.error("Failed to create carousel brief: %s", e)
